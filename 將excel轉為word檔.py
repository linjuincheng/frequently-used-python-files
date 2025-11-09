import os
import glob
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def ensure_rPr(style_element):
    """确保 style_element 下有 <w:rPr>，并返回它。"""
    rPr = style_element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style_element.append(rPr)
    return rPr

def ensure_rFonts(rPr):
    """确保 rPr 下有 <w:rFonts>，并返回它。"""
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    return rFonts

def excel_to_word(excel_path: str, word_path: str, sheet_name: str = None):
    wb = load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    doc = Document()
    # —— 在 Normal 样式里添加 eastAsia 字体 —— 
    normal_style = doc.styles['Normal']
    style_el = normal_style.element
    rPr = ensure_rPr(style_el)
    rFonts = ensure_rFonts(rPr)
    rFonts.set(qn('w:eastAsia'), '新細明體')

    max_row, max_col = ws.max_row, ws.max_column
    table = doc.add_table(rows=max_row, cols=max_col)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row in enumerate(ws.iter_rows()):
        for j, cell in enumerate(row):
            text = '' if cell.value is None else str(cell.value)
            wd_cell = table.cell(i, j)
            p = wd_cell.paragraphs[0]
            run = p.add_run(text)

            # 保留字型大小與粗體
            if cell.font:
                if cell.font.sz:
                    run.font.size = Pt(cell.font.sz)
                if cell.font.bold:
                    run.font.bold = True

    doc.save(word_path)

def batch_convert(folder_path: str):
    patterns = ['*.xlsx', '*.xlsm']
    for pat in patterns:
        for excel_file in glob.glob(os.path.join(folder_path, pat)):
            base = os.path.splitext(os.path.basename(excel_file))[0]
            word_file = os.path.join(folder_path, f"{base}.docx")
            print(f"Converting {excel_file} → {word_file}")
            excel_to_word(excel_file, word_file)
    print("批次轉檔完成。")

if __name__ == "__main__":
    folder = r"C:\_python 2024\2024final"
    batch_convert(folder)
    
    
    
    -------------------------------------------------------------------------------------------
    
    # 多個excel檔(其預設路徑放在C:\_python 2024\2024final), 要將excel檔轉為word檔,並保留excel內的表格格式, 用Python
    
    
    -----------------------------------------------------------------------------------------------------
    
    
